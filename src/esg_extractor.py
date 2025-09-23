#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ESG報告書提取器核心模組 v2.0 - 增強版
支持新關鍵字配置和Word文檔輸出
"""

import json
import re
import os
import sys
import pandas as pd
from pathlib import Path
from typing import Dict, List, Union, Tuple, Optional
from dataclasses import dataclass, asdict
from datetime import datetime
from tqdm import tqdm
import numpy as np

# Word文檔支持
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document as LangchainDocument
from langchain_google_genai import ChatGoogleGenerativeAI

# 添加當前目錄到路徑
sys.path.append(str(Path(__file__).parent))
from config import *
from api_manager import create_api_manager

# =============================================================================
# 增強版關鍵字配置 - 支持新關鍵字
# =============================================================================

class EnhancedKeywordConfig:
    """增強版ESG報告書關鍵字配置 - 支持新的關鍵字組"""
    
    def __init__(self):
        # 原有的再生塑膠關鍵字（保持高準確度）
        self.RECYCLED_PLASTIC_KEYWORDS = {
            "high_relevance_continuous": [
                "再生塑膠", "再生塑料", "再生料", "再生PET", "再生PP",
                "回收塑膠", "回收塑料", "回收PP", "回收PET", 
                "rPET", "PCR塑膠", "PCR塑料", "PCR材料",
                "寶特瓶回收", "廢塑膠回收", "塑膠循環",
                "回收造粒", "再生聚酯", "回收聚酯",
                "循環經濟", "物料回收", "材料回收",
                # 新增關鍵字
                "再生塑膠使用比率", "再生塑膠的使用量", "再生塑膠成本"
            ],
            
            "medium_relevance_continuous": [
                "環保塑膠", "綠色材料", "永續材料",
                "廢料回收", "資源回收", "循環利用",
                # 新增關鍵字
                "材料循環率", "材料可回收率", "再生材料替代率",
                "再生材料使用量", "材料總使用量", "塑膠使用量", "材料使用量"
            ],
            
            "high_relevance_discontinuous": [
                ("寶特瓶", "回收"), ("寶特瓶", "再造"), ("寶特瓶", "循環"),
                ("億支", "寶特瓶"), ("萬支", "寶特瓶"),
                ("PET", "回收"), ("PET", "再生"), ("PP", "回收"), ("PP", "再生"),
                ("塑膠", "回收"), ("塑料", "回收"), ("塑膠", "循環"),
                ("回收", "造粒"), ("回收", "產能"), ("回收", "材料"),
                ("再生", "材料"), ("廢料", "回收"), ("MLCC", "回收"),
                ("原生", "材料"), ("碳排放", "減少"), ("減碳", "效益"),
                ("歷年", "回收"), ("回收", "數量"), ("回收", "效益"),
                ("循環", "經濟"), ("永續", "發展"), ("環保", "產品"),
                # 新增不連續關鍵字組合
                ("材料", "循環"), ("材料", "回收"), ("再生", "使用量"),
                ("碳排", "減量"), ("材質", "分離"), ("單一", "材料")
            ]
        }
        
        # 新增：擴展的環保和永續關鍵字
        self.SUSTAINABILITY_KEYWORDS = {
            "high_relevance_continuous": [
                "再生能源使用率", "綠電憑證", "太陽能電力", "購電協議",
                "再生能源", "再生材料碳排減量", "碳排減量比率", 
                "單位經濟效益", "分選辨視", "成本增加", "碳排放", "材質分離"
            ],
            
            "medium_relevance_continuous": [
                "綠色供應鏈", "永續採購", "環保認證", "低碳製程",
                "節能減碳", "溫室氣體", "碳中和", "淨零排放"
            ],
            
            "high_relevance_discontinuous": [
                ("再生", "能源"), ("太陽能", "發電"), ("綠電", "使用"),
                ("碳排", "減量"), ("溫室", "氣體"), ("環保", "效益"),
                ("永續", "發展"), ("循環", "利用"), ("綠色", "製程")
            ]
        }
        
        # 合併所有關鍵字
        self.ALL_KEYWORDS = {
            "plastic_recycling": self.RECYCLED_PLASTIC_KEYWORDS,
            "sustainability": self.SUSTAINABILITY_KEYWORDS
        }
    
    # 排除規則 - 更加精確
    EXCLUSION_RULES = {
        "exclude_topics": [
            # 職業安全
            "職業災害", "工安", "安全事故", "職災", "工傷", "意外事故",
            
            # 體育賽事
            "馬拉松", "賽事", "選手", "比賽", "賽衣", "運動", "競賽", "體育",
            
            # 水資源（非材料相關）
            "雨水回收", "廢水處理", "水質監測", "用水量", "節水", "水資源",
            
            # 一般改善案（非材料相關）
            "改善案", "改善專案", "案例選拔", "績效改善", "管理改善",
            
            # 能源設備（非材料相關）
            "鍋爐改善", "天然氣燃燒", "燃油改燃", "設備更新",
            
            # 建築產品（非重點）
            "隔熱漆", "節能窗", "隔熱紙", "酷樂漆", "氣密窗", "建材產品"
        ],
        
        "exclude_contexts": [
            # 具體排除上下文
            "垂直馬拉松", "史上最環保賽衣", "各界好手", "運動賽事",
            "職業災害比率", "工安統計", "安全指標",
            "節能改善案", "節水改善案", "優良案例", "改善專案",
            "雨水回收量減少", "降雨量減少", "廢水回收",
            "燃油改燃汽鍋爐", "天然氣燃燒機", "鍋爐改造",
            "酷樂漆", "隔熱漆", "節能氣密窗", "冰酷隔熱紙",
            "員工訓練", "教育訓練", "會議場次", "活動參與"
        ],
        
        "exclude_number_patterns": [
            # 排除特定數值模式
            r'職業災害.*?\d+(?:\.\d+)?(?:%|％|件|次)',
            r'工安.*?\d+(?:\.\d+)?(?:件|次|小時)',
            r'訓練.*?\d+(?:\.\d+)?(?:小時|人次|場)',
            r'會議.*?\d+(?:\.\d+)?(?:次|場|小時)',
            r'降雨量.*?\d+(?:\.\d+)?(?:%|％|mm)',
            r'用水量.*?\d+(?:\.\d+)?(?:噸|立方公尺)'
        ]
    }
    
    # 材料和塑膠相關指標詞 - 更精確
    MATERIAL_INDICATORS = {
        "plastic_materials": [
            "塑膠", "塑料", "聚酯", "PET", "PP", "PE", "PS", "PVC",
            "樹脂", "聚合物", "塑膠粒", "聚酯粒", "塑膠材料",
            "寶特瓶", "瓶片", "容器", "包裝材", "膜材", "纖維材料"
        ],
        
        "recycling_process": [
            "回收", "再生", "循環", "再利用", "回收利用", "資源化",
            "造粒", "再製", "轉換", "處理", "分解", "純化",
            "循環經濟", "廢料處理", "資源循環"
        ],
        
        "production_metrics": [
            "產能", "產量", "生產", "製造", "使用量", "消耗量",
            "替代率", "使用率", "回收率", "循環率", "比例", "比率"
        ]
    }
    
    @classmethod
    def get_all_keywords(cls) -> List[Union[str, tuple]]:
        """獲取所有關鍵字"""
        instance = cls()
        all_keywords = []
        
        # 塑膠回收關鍵字
        all_keywords.extend(instance.RECYCLED_PLASTIC_KEYWORDS["high_relevance_continuous"])
        all_keywords.extend(instance.RECYCLED_PLASTIC_KEYWORDS["medium_relevance_continuous"])
        all_keywords.extend(instance.RECYCLED_PLASTIC_KEYWORDS["high_relevance_discontinuous"])
        
        # 永續發展關鍵字
        all_keywords.extend(instance.SUSTAINABILITY_KEYWORDS["high_relevance_continuous"])
        all_keywords.extend(instance.SUSTAINABILITY_KEYWORDS["medium_relevance_continuous"])
        all_keywords.extend(instance.SUSTAINABILITY_KEYWORDS["high_relevance_discontinuous"])
        
        return all_keywords

# =============================================================================
# 增強版匹配引擎 - 提高準確度
# =============================================================================

class EnhancedESGMatcher:
    """增強版ESG數據匹配引擎 - 提高提取準確度"""
    
    def __init__(self):
        self.config = EnhancedKeywordConfig()
        self.max_distance = 200  # 減少最大距離以提高準確度
        
        # 增強的數值匹配模式
        self.number_patterns = [
            # 基本數量單位
            r'\d+(?:\.\d+)?\s*億支',
            r'\d+(?:\.\d+)?\s*萬支',
            r'\d+(?:,\d{3})*(?:\.\d+)?\s*(?:萬|千)?噸',
            r'\d+(?:,\d{3})*(?:\.\d+)?\s*(?:kg|KG|公斤)',
            
            # 時間相關產能
            r'\d+(?:,\d{3})*(?:\.\d+)?\s*噸/月',
            r'\d+(?:,\d{3})*(?:\.\d+)?\s*噸/年',
            r'\d+(?:,\d{3})*(?:\.\d+)?\s*噸/日',
            
            # 通用計數單位
            r'\d+(?:,\d{3})*(?:\.\d+)?\s*(?:件|個|批|台|套|項)',
            
            # 金額相關
            r'\d+(?:,\d{3})*(?:\.\d+)?\s*(?:億|萬|千)?(?:元|美金|USD)',
            
            # 新增：比率相關數值
            r'\d+(?:\.\d+)?\s*倍',
            r'\d+(?:\.\d+)?\s*(?:ppm|PPM)',
        ]
        
        # 增強的百分比匹配模式
        self.percentage_patterns = [
            r'\d+(?:\.\d+)?\s*%',
            r'\d+(?:\.\d+)?\s*％',
            r'百分之\d+(?:\.\d+)?',
            r'\d+(?:\.\d+)?\s*percent',
        ]
    
    def extract_keyword_value_pairs(self, text: str, keyword: Union[str, tuple]) -> List[Tuple[str, str, float, int]]:
        """提取關鍵字與數值的配對 - 增強版"""
        text_lower = text.lower()
        
        # 1. 檢查關鍵字是否存在
        keyword_match, keyword_confidence, _ = self._match_keyword(text, keyword)
        if not keyword_match:
            return []
        
        # 2. 找到關鍵字位置
        keyword_positions = self._get_keyword_positions(text_lower, keyword)
        if not keyword_positions:
            return []
        
        # 3. 在關鍵字附近尋找數值
        valid_pairs = []
        
        for kw_start, kw_end in keyword_positions:
            # 減小搜索範圍以提高準確度
            search_start = max(0, kw_start - 80)
            search_end = min(len(text), kw_end + 80)
            search_window = text[search_start:search_end]
            
            # 提取數值
            numbers = self._extract_numbers_in_window(search_window)
            percentages = self._extract_percentages_in_window(search_window)
            
            # 驗證數值關聯性（更嚴格的驗證）
            for number in numbers:
                number_pos = search_window.find(number)
                if number_pos != -1:
                    actual_number_pos = search_start + number_pos
                    distance = min(abs(actual_number_pos - kw_start), abs(actual_number_pos - kw_end))
                    
                    association_score = self._calculate_association(
                        text, keyword, number, kw_start, kw_end, actual_number_pos
                    )
                    
                    # 提高關聯度閾值
                    if association_score > 0.6 and distance <= 60:
                        valid_pairs.append((number, 'number', association_score, distance))
            
            for percentage in percentages:
                percentage_pos = search_window.find(percentage)
                if percentage_pos != -1:
                    actual_percentage_pos = search_start + percentage_pos
                    distance = min(abs(actual_percentage_pos - kw_start), abs(actual_percentage_pos - kw_end))
                    
                    association_score = self._calculate_association(
                        text, keyword, percentage, kw_start, kw_end, actual_percentage_pos
                    )
                    
                    # 提高關聯度閾值
                    if association_score > 0.6 and distance <= 60:
                        valid_pairs.append((percentage, 'percentage', association_score, distance))
        
        # 去重並排序，只保留最佳結果
        unique_pairs = []
        seen_values = set()
        
        for value, value_type, score, distance in sorted(valid_pairs, key=lambda x: (-x[2], x[3])):
            if value not in seen_values:
                seen_values.add(value)
                unique_pairs.append((value, value_type, score, distance))
        
        return unique_pairs[:2]  # 只保留前2個最佳匹配
    
    def comprehensive_relevance_check(self, text: str, keyword: Union[str, tuple]) -> Tuple[bool, float, str]:
        """綜合相關性檢查 - 增強版"""
        text_lower = text.lower()
        
        # 1. 強排除檢查 - 更嚴格
        exclusion_score = self._check_strong_exclusions(text_lower)
        if exclusion_score > 0.3:  # 降低排除閾值，更容易排除
            return False, 0.0, f"排除內容: {exclusion_score:.2f}"
        
        # 2. 關鍵字匹配
        keyword_match, keyword_confidence, keyword_details = self._match_keyword(text, keyword)
        if not keyword_match:
            return False, 0.0, "關鍵字不匹配"
        
        # 3. 材料相關性檢查 - 更嚴格
        material_relevance = self._check_material_relevance(text_lower)
        if material_relevance < 0.4:  # 提高材料相關性閾值
            return False, 0.0, f"非材料相關: {material_relevance:.2f}"
        
        # 4. 上下文質量檢查
        context_quality = self._check_context_quality(text_lower)
        if context_quality < 0.3:
            return False, 0.0, f"上下文質量不足: {context_quality:.2f}"
        
        # 5. 計算綜合分數
        final_score = (
            keyword_confidence * 0.25 + 
            material_relevance * 0.30 + 
            context_quality * 0.25 +
            (1 - exclusion_score) * 0.20  # 排除懲罰
        )
        
        # 提高相關性閾值
        is_relevant = final_score > 0.65
        
        details = f"關鍵字:{keyword_confidence:.2f}, 材料:{material_relevance:.2f}, 上下文:{context_quality:.2f}, 排除:{exclusion_score:.2f}"
        
        return is_relevant, final_score, details
    
    # 新增輔助方法
    def _check_strong_exclusions(self, text: str) -> float:
        """檢查強排除模式 - 增強版"""
        exclusion_score = 0.0
        
        # 檢查排除主題（權重更高）
        for topic in self.config.EXCLUSION_RULES["exclude_topics"]:
            if topic in text:
                exclusion_score += 0.25
        
        # 檢查排除上下文（權重最高）
        for context in self.config.EXCLUSION_RULES["exclude_contexts"]:
            if context in text:
                exclusion_score += 0.35
        
        # 檢查排除數值模式
        for pattern in self.config.EXCLUSION_RULES["exclude_number_patterns"]:
            if re.search(pattern, text, re.IGNORECASE):
                exclusion_score += 0.30
        
        return min(exclusion_score, 1.0)
    
    def _check_material_relevance(self, text: str) -> float:
        """檢查材料相關性 - 更精確"""
        material_score = 0.0
        recycling_score = 0.0
        production_score = 0.0
        
        # 材料相關詞計分
        material_count = 0
        for indicator in self.config.MATERIAL_INDICATORS["plastic_materials"]:
            if indicator in text:
                material_count += 1
        material_score = min(material_count / 3.0, 1.0)
        
        # 回收處理相關詞計分
        recycling_count = 0
        for indicator in self.config.MATERIAL_INDICATORS["recycling_process"]:
            if indicator in text:
                recycling_count += 1
        recycling_score = min(recycling_count / 2.0, 1.0)
        
        # 生產指標相關詞計分
        production_count = 0
        for indicator in self.config.MATERIAL_INDICATORS["production_metrics"]:
            if indicator in text:
                production_count += 1
        production_score = min(production_count / 2.0, 1.0)
        
        # 必須同時具備材料和處理/生產相關詞
        if material_score > 0 and (recycling_score > 0 or production_score > 0):
            return (material_score + recycling_score + production_score) / 3.0
        else:
            return 0.0
    
    def _check_context_quality(self, text: str) -> float:
        """檢查上下文質量"""
        quality_indicators = [
            "使用", "生產", "製造", "應用", "處理", "回收",
            "數量", "比例", "比率", "產能", "效益", "成本",
            "減少", "增加", "提高", "降低", "改善", "優化"
        ]
        
        found_indicators = sum(1 for indicator in quality_indicators if indicator in text)
        quality_score = min(found_indicators / 4.0, 1.0)
        
        # 檢查數值相關性
        has_meaningful_numbers = bool(re.search(r'\d+(?:,\d{3})*(?:\.\d+)?\s*(?:噸|億|萬|%|％)', text))
        if has_meaningful_numbers:
            quality_score += 0.3
        
        return min(quality_score, 1.0)
    
    # 保留原有的輔助方法（簡化顯示）
    def _match_keyword(self, text: str, keyword: Union[str, tuple]) -> Tuple[bool, float, str]:
        """關鍵字匹配 - 保持原邏輯"""
        text_lower = text.lower()
        
        if isinstance(keyword, str):
            if keyword.lower() in text_lower:
                return True, 1.0, f"精確匹配: {keyword}"
            return False, 0.0, ""
        
        elif isinstance(keyword, tuple):
            components = [comp.lower() for comp in keyword]
            positions = []
            
            for comp in components:
                pos = text_lower.find(comp)
                if pos == -1:
                    return False, 0.0, f"缺少組件: {comp}"
                positions.append(pos)
            
            distance = max(positions) - min(positions)
            
            if distance <= 50:
                return True, 0.9, f"近距離匹配({distance}字)"
            elif distance <= 100:
                return True, 0.8, f"中距離匹配({distance}字)"
            elif distance <= self.max_distance:
                return True, 0.6, f"遠距離匹配({distance}字)"
            else:
                return True, 0.4, f"極遠距離匹配({distance}字)"
        
        return False, 0.0, ""
    
    def _get_keyword_positions(self, text: str, keyword: Union[str, tuple]) -> List[Tuple[int, int]]:
        """獲取關鍵字位置"""
        positions = []
        
        if isinstance(keyword, str):
            keyword_lower = keyword.lower()
            start = 0
            while True:
                pos = text.find(keyword_lower, start)
                if pos == -1:
                    break
                positions.append((pos, pos + len(keyword_lower)))
                start = pos + 1
        
        elif isinstance(keyword, tuple):
            components = [comp.lower() for comp in keyword]
            component_positions = {}
            
            for comp in components:
                comp_positions = []
                start = 0
                while True:
                    pos = text.find(comp, start)
                    if pos == -1:
                        break
                    comp_positions.append((pos, pos + len(comp)))
                    start = pos + 1
                component_positions[comp] = comp_positions
            
            for comp1_pos in component_positions.get(components[0], []):
                for comp2_pos in component_positions.get(components[1], []):
                    distance = abs(comp1_pos[0] - comp2_pos[0])
                    if distance <= self.max_distance:
                        start_pos = min(comp1_pos[0], comp2_pos[0])
                        end_pos = max(comp1_pos[1], comp2_pos[1])
                        positions.append((start_pos, end_pos))
        
        return positions
    
    def _extract_numbers_in_window(self, window_text: str) -> List[str]:
        """提取數值"""
        numbers = []
        for pattern in self.number_patterns:
            matches = re.findall(pattern, window_text, re.IGNORECASE)
            numbers.extend(matches)
        return list(set(numbers))
    
    def _extract_percentages_in_window(self, window_text: str) -> List[str]:
        """提取百分比"""
        percentages = []
        for pattern in self.percentage_patterns:
            matches = re.findall(pattern, window_text, re.IGNORECASE)
            percentages.extend(matches)
        return list(set(percentages))
    
    def _calculate_association(self, text: str, keyword: Union[str, tuple], 
                             value: str, kw_start: int, kw_end: int, value_pos: int) -> float:
        """計算關鍵字與數值的關聯度 - 增強版"""
        
        # 距離因子（更嚴格）
        distance = min(abs(value_pos - kw_start), abs(value_pos - kw_end))
        if distance <= 15:
            distance_score = 1.0
        elif distance <= 30:
            distance_score = 0.8
        elif distance <= 60:
            distance_score = 0.6
        else:
            distance_score = 0.3
        
        # 上下文相關性
        context_start = min(kw_start, value_pos) - 25
        context_end = max(kw_end, value_pos + len(value)) + 25
        context_start = max(0, context_start)
        context_end = min(len(text), context_end)
        context = text[context_start:context_end].lower()
        
        context_score = self._calculate_context_score(context)
        
        # 數值合理性
        value_score = self._calculate_value_score(value, context)
        
        # 語義連接詞檢查
        connection_score = self._check_semantic_connection(context, keyword)
        
        final_score = (
            distance_score * 0.30 +
            context_score * 0.30 + 
            value_score * 0.25 +
            connection_score * 0.15
        )
        
        return final_score
    
    def _check_semantic_connection(self, context: str, keyword: Union[str, tuple]) -> float:
        """檢查語義連接詞"""
        connection_words = [
            "達到", "為", "約", "共", "總計", "累計", "實現", "完成",
            "使用", "生產", "製造", "處理", "回收", "節省", "減少", "增加"
        ]
        
        connection_score = 0.0
        for word in connection_words:
            if word in context:
                connection_score += 0.2
        
        return min(connection_score, 1.0)
    
    def _calculate_context_score(self, context: str) -> float:
        """計算上下文分數 - 增強版"""
        high_relevance_words = [
            "回收", "再生", "循環", "製造", "生產", "產能", "使用",
            "塑膠", "塑料", "聚酯", "材料", "寶特瓶", "減碳", "效益",
            "比例", "比率", "數量", "產量", "使用量"
        ]
        
        negative_words = [
            "災害", "事故", "馬拉松", "賽事", "改善案", "案例",
            "雨水", "節能", "隔熱", "鍋爐", "燃油", "訓練", "會議"
        ]
        
        score = 0.0
        
        for word in high_relevance_words:
            if word in context:
                score += 0.12
        
        for word in negative_words:
            if word in context:
                score -= 0.25
        
        return max(0.0, min(1.0, score))
    
    def _calculate_value_score(self, value: str, context: str) -> float:
        """計算數值合理性分數 - 增強版"""
        number_match = re.search(r'\d+(?:,\d{3})*(?:\.\d+)?', value)
        if not number_match:
            return 0.0
        
        try:
            number_str = number_match.group().replace(',', '')
            number = float(number_str)
        except ValueError:
            return 0.0
        
        # 根據單位評估合理性
        if "億支" in value:
            if 0.5 <= number <= 100:
                return 1.0
            elif 0.1 <= number <= 300:
                return 0.7
            else:
                return 0.3
        
        elif "萬噸" in value or "千噸" in value:
            if 0.1 <= number <= 50:
                return 1.0
            elif 0.01 <= number <= 200:
                return 0.7
            else:
                return 0.3
        
        elif "噸" in value and "萬" not in value and "千" not in value:
            if 10 <= number <= 50000:
                return 1.0
            elif 1 <= number <= 100000:
                return 0.7
            else:
                return 0.3
        
        elif "%" in value or "％" in value:
            if 0.1 <= number <= 100:
                return 1.0
            else:
                return 0.2
        
        elif "件" in value:
            if 10 <= number <= 10000:
                return 1.0
            elif 1 <= number <= 50000:
                return 0.7
            else:
                return 0.3
        
        return 0.5

# =============================================================================
# Word文檔輸出功能
# =============================================================================

class ESGWordExporter:
    """ESG Word文檔導出器"""
    
    def __init__(self):
        pass
    
    def create_word_document(self, extractions: List, doc_info, stock_code: str, short_company_name: str) -> str:
        """創建Word文檔"""
        
        # 生成Word檔案名
        if stock_code:
            word_filename = f"提取統整_{stock_code}_{short_company_name}_{doc_info.report_year}.docx"
        else:
            company_safe = re.sub(r'[^\w\s-]', '', short_company_name).strip()
            word_filename = f"提取統整_{company_safe}_{doc_info.report_year}.docx"
        
        word_path = os.path.join(RESULTS_PATH, word_filename)
        
        # 創建Word文檔
        doc = Document()
        
        # 設置文檔標題
        title = doc.add_heading(f'{short_company_name} {doc_info.report_year}年 ESG提取統整報告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加摘要信息
        summary_para = doc.add_paragraph()
        summary_para.add_run(f'股票代號: ').bold = True
        summary_para.add_run(f'{stock_code or "N/A"}\n')
        summary_para.add_run(f'公司全稱: ').bold = True
        summary_para.add_run(f'{doc_info.company_name}\n')
        summary_para.add_run(f'報告年度: ').bold = True
        summary_para.add_run(f'{doc_info.report_year}\n')
        summary_para.add_run(f'提取時間: ').bold = True
        summary_para.add_run(f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        summary_para.add_run(f'提取結果數量: ').bold = True
        summary_para.add_run(f'{len(extractions)} 項')
        
        # 添加分隔線
        doc.add_paragraph('=' * 80)
        
        if not extractions:
            doc.add_heading('未找到相關提取結果', level=1)
            no_result_para = doc.add_paragraph()
            no_result_para.add_run('在此份ESG報告中未找到再生塑膠或永續材料相關的數值數據。\n\n')
            no_result_para.add_run('可能的原因：\n')
            no_result_para.add_run('1. 該公司未涉及相關業務\n')
            no_result_para.add_run('2. 報告中未詳細披露相關數據\n')
            no_result_para.add_run('3. 關鍵字匹配範圍需要調整')
        else:
            # 按頁碼排序
            sorted_extractions = sorted(extractions, key=lambda x: (
                int(re.findall(r'\d+', x.page_number)[0]) if re.findall(r'\d+', x.page_number) else 999,
                x.confidence
            ), reverse=False)
            
            for i, extraction in enumerate(sorted_extractions, 1):
                # 添加序號標題
                doc.add_heading(f'提取結果 {i}', level=2)
                
                # 頁碼
                page_para = doc.add_paragraph()
                page_para.add_run('頁碼: ').bold = True
                page_para.add_run(f'{extraction.page_number}')
                
                # 關鍵字
                keyword_para = doc.add_paragraph()
                keyword_para.add_run('關鍵字: ').bold = True
                keyword_para.add_run(f'{extraction.keyword}')
                
                # 數值
                value_para = doc.add_paragraph()
                value_para.add_run('數值: ').bold = True
                if extraction.value == "[相關描述]":
                    value_para.add_run('相關描述（無具體數值）')
                else:
                    value_para.add_run(f'{extraction.value}')
                    if extraction.unit:
                        value_para.add_run(f' ({extraction.unit})')
                
                # 信心分數
                confidence_para = doc.add_paragraph()
                confidence_para.add_run('信心分數: ').bold = True
                confidence_para.add_run(f'{extraction.confidence:.3f}')
                
                # 整個段落內容
                content_para = doc.add_paragraph()
                content_para.add_run('段落內容: ').bold = True
                content_para.add_run('\n')
                
                # 段落內容使用不同的字體
                content_run = content_para.add_run(extraction.paragraph)
                content_run.font.name = '微軟正黑體'
                content_run.font.size = Pt(10)
                
                # 添加分隔線
                if i < len(sorted_extractions):
                    doc.add_paragraph('-' * 60)
        
        # 添加頁尾信息
        footer_para = doc.add_paragraph()
        footer_para.add_run('\n\n生成工具: ').italic = True
        footer_para.add_run('ESG報告書提取器 v2.0').italic = True
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 保存文檔
        doc.save(word_path)
        
        return word_path

# =============================================================================
# 股票代號映射器（保持不變）
# =============================================================================

class StockCodeMapper:
    """股票代號與公司名稱對照映射器"""
    
    def __init__(self):
        # 台灣主要上市公司股票代號對照表
        self.stock_code_mapping = {
            # 石化塑膠類
            "1301": "台塑",
            "1303": "南亞", 
            "1326": "台化",
            "1314": "中石化",
            "1605": "華新",
            "1722": "台肥",
            "6505": "台塑化",
            
            # 電子科技類
            "2330": "台積電",
            "2454": "聯發科",
            "2317": "鴻海",
            "2382": "廣達",
            "2308": "台達電",
            "3045": "台灣大",
            "4904": "遠傳",
            "3008": "大立光",
            "2474": "可成",
            "6770": "力積電",
            
            # 金融類
            "2884": "玉山金",
            "2885": "元大金", 
            "2886": "兆豐金",
            "2887": "台新金",
            "2888": "新光金",
            "2891": "中信金",
            "2892": "第一金",
            "2880": "華南金",
            
            # 傳統產業
            "2002": "中鋼",
            "1216": "統一",
            "1101": "台泥",
            "2408": "南科",
            "2409": "友達",
            "2412": "中華電",
            "3481": "群創",
            "2207": "和泰車",
            "2912": "統一超",
            "6278": "台表科",
        }
        
        # 公司名稱反向對照（用於查找股票代號）
        self.reverse_mapping = {}
        for code, name in self.stock_code_mapping.items():
            self.reverse_mapping[name] = code
            # 添加完整公司名稱的對照
            full_names = {
                "台塑": ["台灣塑膠工業", "台塑工業", "台灣塑膠", "台塑公司"],
                "南亞": ["南亞塑膠工業", "南亞塑膠", "南亞公司"],
                "台化": ["台灣化學纖維", "台化公司", "台灣化纖"],
                "台積電": ["台灣積體電路", "台灣積體電路製造", "TSMC"],
                "鴻海": ["鴻海精密", "鴻海科技"],
                "中鋼": ["中國鋼鐵", "中鋼公司"],
                "台泥": ["台灣水泥", "台泥公司"],
                "統一": ["統一企業", "統一公司"],
                "中華電": ["中華電信", "中華電信公司"]
            }
            
            if name in full_names:
                for full_name in full_names[name]:
                    self.reverse_mapping[full_name] = code
    
    def extract_stock_info_from_vector_name(self, vector_name: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
        """
        從向量資料庫名稱中提取股票代號、公司名稱、年度
        例如：esg_db_1301_台塑_2024_esg報告書 -> ("1301", "台塑", "2024")
        """
        try:
            # 移除前綴和後綴
            cleaned_name = vector_name.replace("esg_db_", "")
            
            # 嘗試匹配格式：股票代號_公司名稱_年度_其他
            pattern = r'^(\d{4})_([^_]+)_(\d{4})(?:_.*)?$'
            match = re.match(pattern, cleaned_name)
            
            if match:
                stock_code = match.group(1)
                company_name = match.group(2)
                year = match.group(3)
                return stock_code, company_name, year
            
            # 如果沒有股票代號，嘗試匹配：公司名稱_年度
            pattern2 = r'^([^_]+)_(\d{4})(?:_.*)?$'
            match2 = re.match(pattern2, cleaned_name)
            
            if match2:
                company_name = match2.group(1)
                year = match2.group(2)
                # 嘗試查找股票代號
                stock_code = self.find_stock_code_by_company(company_name)
                return stock_code, company_name, year
            
            return None, None, None
            
        except Exception as e:
            print(f"⚠️ 解析向量名稱失敗: {e}")
            return None, None, None
    
    def find_stock_code_by_company(self, company_name: str) -> Optional[str]:
        """根據公司名稱查找股票代號"""
        if not company_name:
            return None
        
        # 直接匹配
        if company_name in self.reverse_mapping:
            return self.reverse_mapping[company_name]
        
        # 模糊匹配
        company_clean = company_name.replace("股份有限公司", "").replace("有限公司", "").replace("公司", "").strip()
        
        for mapped_name, code in self.reverse_mapping.items():
            if company_clean in mapped_name or mapped_name in company_clean:
                return code
        
        # 關鍵字匹配
        for mapped_name, code in self.reverse_mapping.items():
            if len(mapped_name) >= 2 and mapped_name in company_name:
                return code
        
        return None
    
    def get_short_company_name(self, company_name: str, stock_code: str = None) -> str:
        """獲取簡化的公司名稱"""
        if not company_name:
            return "未知公司"
        
        # 如果有股票代號，直接使用對照表中的簡稱
        if stock_code and stock_code in self.stock_code_mapping:
            return self.stock_code_mapping[stock_code]
        
        # 手動簡化常見公司名稱
        simplifications = {
            "台灣塑膠工業股份有限公司": "台塑",
            "台灣塑膠工業": "台塑", 
            "台塑工業": "台塑",
            "南亞塑膠工業股份有限公司": "南亞",
            "南亞塑膠工業": "南亞",
            "南亞塑膠": "南亞",
            "台灣化學纖維股份有限公司": "台化",
            "台灣化學纖維": "台化",
            "台化公司": "台化",
            "台灣積體電路製造股份有限公司": "台積電",
            "台灣積體電路": "台積電",
            "鴻海精密工業股份有限公司": "鴻海",
            "鴻海精密": "鴻海",
            "中國鋼鐵股份有限公司": "中鋼",
            "中國鋼鐵": "中鋼",
            "台灣水泥股份有限公司": "台泥", 
            "台灣水泥": "台泥",
            "統一企業股份有限公司": "統一",
            "統一企業": "統一",
            "中華電信股份有限公司": "中華電",
            "中華電信": "中華電"
        }
        
        # 精確匹配
        if company_name in simplifications:
            return simplifications[company_name]
        
        # 移除常見後綴並簡化
        simplified = company_name
        suffixes = ["股份有限公司", "有限公司", "股份公司", "公司", "集團", "企業"]
        
        for suffix in suffixes:
            if simplified.endswith(suffix):
                simplified = simplified[:-len(suffix)].strip()
                break
        
        # 如果簡化後長度合適，返回簡化結果
        if 2 <= len(simplified) <= 4:
            return simplified
        
        # 否則返回前4個字符
        return simplified[:4] if len(simplified) > 4 else simplified

# =============================================================================
# 原有數據結構（保持不變）
# =============================================================================

@dataclass
class DocumentInfo:
    """文檔信息"""
    company_name: str
    report_year: str
    pdf_name: str
    db_path: str

@dataclass
class NumericExtraction:
    """數值提取結果"""
    keyword: str
    value: str
    value_type: str
    unit: str
    paragraph: str
    paragraph_number: int
    page_number: str
    confidence: float
    context_window: str
    company_name: str = ""
    report_year: str = ""
    keyword_distance: int = 0

@dataclass
class ProcessingSummary:
    """處理摘要"""
    company_name: str
    report_year: str
    total_documents: int
    stage1_passed: int
    stage2_passed: int
    total_extractions: int
    keywords_found: Dict[str, int]
    processing_time: float

# =============================================================================
# 增強版ESG提取器主類
# =============================================================================

class EnhancedESGExtractor:
    """增強版ESG報告書提取器主類"""
    
    def __init__(self, enable_llm: bool = True):
        self.enable_llm = enable_llm
        self.matcher = EnhancedESGMatcher()  # 使用增強版匹配器
        self.keyword_config = EnhancedKeywordConfig()  # 使用增強版關鍵字配置
        self.stock_mapper = StockCodeMapper()
        self.word_exporter = ESGWordExporter()  # 新增Word導出器
        
        if self.enable_llm:
            self._init_llm()
        
        print("✅ 增強版ESG報告書提取器初始化完成")
        print("🔧 新功能：支援擴展關鍵字、提高提取準確度、Word文檔輸出")

    def _init_llm(self):
        """初始化LLM"""
        try:
            print("🤖 初始化Gemini API管理器...")
            self.api_manager = create_api_manager()
            print("✅ LLM初始化完成")
        except Exception as e:
            print(f"⚠️ LLM初始化失敗: {e}")
            self.enable_llm = False
    
    def process_single_document(self, doc_info: DocumentInfo, max_documents: int = 400) -> Tuple[List[NumericExtraction], ProcessingSummary, str, str]:
        """處理單個文檔 - 增強版，返回Excel和Word檔案路徑"""
        start_time = datetime.now()
        print(f"\n📊 處理文檔: {doc_info.company_name} - {doc_info.report_year}")
        print("=" * 60)
        
        # 1. 載入向量資料庫
        db = self._load_vector_database(doc_info.db_path)
        
        # 2. 文檔檢索
        documents = self._document_retrieval(db, max_documents)
        
        # 3. 數據提取（使用增強版匹配器）
        extractions = self._extract_data(documents, doc_info)
        
        # 4. 後處理
        extractions = self._post_process_extractions(extractions)
        
        # 5. 創建處理摘要
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        keywords_found = {}
        for extraction in extractions:
            keyword = extraction.keyword
            keywords_found[keyword] = keywords_found.get(keyword, 0) + 1
        
        summary = ProcessingSummary(
            company_name=doc_info.company_name,
            report_year=doc_info.report_year,
            total_documents=len(documents),
            stage1_passed=len(documents),
            stage2_passed=len(extractions),
            total_extractions=len(extractions),
            keywords_found=keywords_found,
            processing_time=processing_time
        )
        
        # 6. 獲取股票信息
        vector_db_name = Path(doc_info.db_path).name
        stock_code, extracted_company, extracted_year = self.stock_mapper.extract_stock_info_from_vector_name(vector_db_name)
        
        final_company = extracted_company if extracted_company else doc_info.company_name
        final_year = extracted_year if extracted_year else doc_info.report_year
        final_stock_code = stock_code
        
        if not final_stock_code:
            final_stock_code = self.stock_mapper.find_stock_code_by_company(final_company)
        
        short_company_name = self.stock_mapper.get_short_company_name(final_company, final_stock_code)
        
        # 7. 匯出Excel結果
        excel_path = self._export_to_excel(extractions, summary, doc_info, final_stock_code, short_company_name)
        
        # 8. 匯出Word結果
        word_path = self.word_exporter.create_word_document(extractions, doc_info, final_stock_code, short_company_name)
        
        print(f"📊 Excel檔案: {Path(excel_path).name}")
        print(f"📝 Word檔案: {Path(word_path).name}")
        
        return extractions, summary, excel_path, word_path
    
    def process_multiple_documents(self, docs_info: Dict[str, DocumentInfo], max_documents: int = 400) -> Dict[str, Tuple]:
        """批量處理多個文檔 - 增強版"""
        print(f"📊 開始批量處理 {len(docs_info)} 個文檔")
        print("=" * 60)
        
        results = {}
        
        for pdf_path, doc_info in docs_info.items():
            try:
                print(f"\n📄 處理: {doc_info.company_name} - {doc_info.report_year}")
                
                extractions, summary, excel_path, word_path = self.process_single_document(doc_info, max_documents)
                
                results[pdf_path] = (extractions, summary, excel_path, word_path)
                
                print(f"✅ 完成: 生成 {len(extractions)} 個結果")
                print(f"   📊 Excel: {Path(excel_path).name}")
                print(f"   📝 Word: {Path(word_path).name}")
                
            except Exception as e:
                print(f"❌ 處理失敗 {doc_info.company_name}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        print(f"\n🎉 批量處理完成！成功處理 {len(results)}/{len(docs_info)} 個文檔")
        return results
    
    # 以下方法大部分保持不變，只修改關鍵部分
    
    def _load_vector_database(self, db_path: str):
        """載入向量資料庫"""
        if not os.path.exists(db_path):
            raise FileNotFoundError(f"向量資料庫不存在: {db_path}")
        
        embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={'device': 'cpu'}
        )
        db = FAISS.load_local(
            db_path, 
            embeddings, 
            allow_dangerous_deserialization=True
        )
        return db
    
    def _document_retrieval(self, db, max_docs: int) -> List[LangchainDocument]:
        """文檔檢索 - 使用新的關鍵字配置"""
        keywords = self.keyword_config.get_all_keywords()
        all_docs = []
        
        # 關鍵字檢索
        print("   🔍 執行關鍵字檢索...")
        for keyword in keywords[:30]:  # 增加檢索範圍
            search_term = keyword if isinstance(keyword, str) else " ".join(keyword)
            docs = db.similarity_search(search_term, k=12)
            all_docs.extend(docs)
        
        # 主題檢索 - 增強版
        print("   🔍 執行主題檢索...")
        topic_queries = [
            "塑膠 回收 材料",
            "寶特瓶 循環 經濟",
            "再生 環保 永續",
            "廢料 處理 利用",
            "材料 循環 使用",  # 新增主題
            "碳排 減量 效益",  # 新增主題
            "再生能源 綠電",   # 新增主題
        ]
        
        for query in topic_queries:
            docs = db.similarity_search(query, k=15)
            all_docs.extend(docs)
        
        # 數值檢索
        print("   🔍 執行數值檢索...")
        number_queries = [
            "億支", "萬噸", "千噸", "產能", "回收量", "使用量",
            "減碳", "百分比", "效益", "數量", "比例", "比率"
        ]
        
        for query in number_queries:
            docs = db.similarity_search(query, k=8)
            all_docs.extend(docs)
        
        # 去重
        unique_docs = {}
        for doc in all_docs:
            doc_hash = hash(doc.page_content)
            if doc_hash not in unique_docs:
                unique_docs[doc_hash] = doc
        
        result_docs = list(unique_docs.values())[:max_docs]
        print(f"📚 檢索到 {len(result_docs)} 個候選文檔")
        return result_docs
    
    def _extract_data(self, documents: List[LangchainDocument], doc_info: DocumentInfo) -> List[NumericExtraction]:
        """數據提取 - 使用增強版匹配器"""
        print("🎯 執行數據提取（增強版）...")
        
        keywords = self.keyword_config.get_all_keywords()
        extractions = []
        
        for doc in tqdm(documents, desc="數據提取"):
            # 段落分割
            paragraphs = self._split_paragraphs(doc.page_content)
            page_num = doc.metadata.get('page', '未知')
            
            for para_idx, paragraph in enumerate(paragraphs):
                if len(paragraph.strip()) < 20:  # 提高最小段落長度
                    continue
                
                # 對每個關鍵字進行匹配
                for keyword in keywords:
                    # 使用增強版相關性檢查
                    is_relevant, relevance_score, details = self.matcher.comprehensive_relevance_check(paragraph, keyword)
                    
                    if is_relevant and relevance_score > 0.65:  # 提高相關性閾值
                        # 提取數值配對
                        value_pairs = self.matcher.extract_keyword_value_pairs(paragraph, keyword)
                        
                        # 如果沒有找到數值但相關性很高，保留作為描述
                        if not value_pairs and relevance_score > 0.80:  # 提高描述保留閾值
                            keyword_str = keyword if isinstance(keyword, str) else " + ".join(keyword)
                            
                            extraction = NumericExtraction(
                                keyword=keyword_str,
                                value="[相關描述]",
                                value_type='description',
                                unit='',
                                paragraph=paragraph.strip(),
                                paragraph_number=para_idx + 1,
                                page_number=f"第{page_num}頁",
                                confidence=relevance_score,
                                context_window=self._get_context_window(doc.page_content, paragraph),
                                company_name=doc_info.company_name,
                                report_year=doc_info.report_year,
                                keyword_distance=0
                            )
                            extractions.append(extraction)
                        
                        # 處理找到的數值配對
                        for value, value_type, association_score, distance in value_pairs:
                            keyword_str = keyword if isinstance(keyword, str) else " + ".join(keyword)
                            
                            final_confidence = (relevance_score * 0.4 + association_score * 0.6)
                            
                            # 只保留高信心度的結果
                            if final_confidence > 0.70:
                                extraction = NumericExtraction(
                                    keyword=keyword_str,
                                    value=value,
                                    value_type=value_type,
                                    unit=self._extract_unit(value) if value_type == 'number' else '%',
                                    paragraph=paragraph.strip(),
                                    paragraph_number=para_idx + 1,
                                    page_number=f"第{page_num}頁",
                                    confidence=final_confidence,
                                    context_window=self._get_context_window(doc.page_content, paragraph),
                                    company_name=doc_info.company_name,
                                    report_year=doc_info.report_year,
                                    keyword_distance=distance
                                )
                                extractions.append(extraction)
        
        print(f"✅ 數據提取完成: 找到 {len(extractions)} 個結果")
        return extractions
    
    def _post_process_extractions(self, extractions: List[NumericExtraction]) -> List[NumericExtraction]:
        """後處理和去重 - 增強版"""
        if not extractions:
            return extractions
        
        print(f"🔧 後處理 {len(extractions)} 個提取結果...")
        
        # 精確去重（更嚴格）
        unique_extractions = []
        seen_combinations = set()
        
        for extraction in extractions:
            identifier = (
                extraction.keyword,
                extraction.value,
                extraction.value_type,
                extraction.paragraph[:150]  # 增加段落檢查長度
            )
            
            if identifier not in seen_combinations:
                seen_combinations.add(identifier)
                unique_extractions.append(extraction)
        
        print(f"📊 精確去重後: {len(unique_extractions)} 個結果")
        
        # 頁面去重（每頁最多保留3筆高質量結果）
        page_filtered_extractions = self._apply_per_page_filtering(unique_extractions, max_per_page=3)
        
        # 按信心分數排序
        page_filtered_extractions.sort(key=lambda x: x.confidence, reverse=True)
        
        print(f"✅ 後處理完成: 保留 {len(page_filtered_extractions)} 個最終結果")
        return page_filtered_extractions
    
    def _apply_per_page_filtering(self, extractions: List[NumericExtraction], max_per_page: int = 3) -> List[NumericExtraction]:
        """按頁面去重"""
        if not extractions:
            return extractions
        
        print(f"📄 執行按頁面去重（每頁最多保留 {max_per_page} 筆）...")
        
        # 按頁碼分組
        page_groups = {}
        for extraction in extractions:
            page_key = str(extraction.page_number).strip()
            if page_key not in page_groups:
                page_groups[page_key] = []
            page_groups[page_key].append(extraction)
        
        # 每頁面內按信心分數排序並保留最佳結果
        filtered_extractions = []
        
        for page_key, page_extractions in page_groups.items():
            # 按信心分數排序
            page_extractions.sort(key=lambda x: x.confidence, reverse=True)
            
            # 進一步去重：避免同一頁面的重複內容
            page_unique = []
            seen_values = set()
            
            for extraction in page_extractions:
                value_key = (extraction.value, extraction.keyword)
                if value_key not in seen_values:
                    seen_values.add(value_key)
                    page_unique.append(extraction)
            
            kept_extractions = page_unique[:max_per_page]
            filtered_extractions.extend(kept_extractions)
        
        print(f"   ✅ 頁面去重完成: {len(filtered_extractions)} 筆最終結果")
        return filtered_extractions
    
    def _export_to_excel(self, extractions: List[NumericExtraction], summary: ProcessingSummary, 
                        doc_info: DocumentInfo, stock_code: str, short_company_name: str) -> str:
        """匯出結果到Excel - 保持原有邏輯"""
        # 生成Excel檔案名
        if stock_code:
            if len(extractions) == 0:
                output_filename = f"無提取_{stock_code}_{short_company_name}_{doc_info.report_year}.xlsx"
                status_message = "無提取結果"
            else:
                output_filename = f"提取結果_{stock_code}_{short_company_name}_{doc_info.report_year}.xlsx"
                status_message = f"提取結果: {len(extractions)} 項"
        else:
            company_safe = re.sub(r'[^\w\s-]', '', short_company_name).strip()
            if len(extractions) == 0:
                output_filename = f"無提取_{company_safe}_{doc_info.report_year}.xlsx"
                status_message = "無提取結果"
            else:
                output_filename = f"提取結果_{company_safe}_{doc_info.report_year}.xlsx"
                status_message = f"提取結果: {len(extractions)} 項"
        
        output_path = os.path.join(RESULTS_PATH, output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        print(f"📊 匯出Excel結果: {output_filename}")
        if stock_code:
            print(f"   🏢 {stock_code} - {short_company_name} ({doc_info.report_year})")
        
        # 準備主要數據
        main_data = []
        
        # 第一行：公司信息（包含股票代號）
        header_row = {
            '關鍵字': f"股票代號: {stock_code or 'N/A'} | 公司: {doc_info.company_name}",
            '提取數值': f"報告年度: {doc_info.report_year}",
            '數據類型': f"處理時間: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            '單位': f"向量庫: {Path(doc_info.db_path).name}",
            '段落內容': f"{status_message}（ESG報告書提取器 v2.0 增強版）",
            '段落編號': '',
            '頁碼': '',
            '信心分數': '',
            '上下文': ''
        }
        main_data.append(header_row)
        
        # 空行分隔
        main_data.append({col: '' for col in header_row.keys()})
        
        # 如果有提取結果，添加結果數據
        if len(extractions) > 0:
            for extraction in extractions:
                main_data.append({
                    '關鍵字': extraction.keyword,
                    '提取數值': extraction.value,
                    '數據類型': extraction.value_type,
                    '單位': extraction.unit,
                    '段落內容': extraction.paragraph,
                    '段落編號': extraction.paragraph_number,
                    '頁碼': extraction.page_number,
                    '信心分數': round(extraction.confidence, 3),
                    '上下文': extraction.context_window[:200] + "..." if len(extraction.context_window) > 200 else extraction.context_window
                })
        else:
            # 如果沒有提取結果，添加說明行
            no_result_row = {
                '關鍵字': '無相關關鍵字匹配',
                '提取數值': 'N/A',
                '數據類型': 'no_data',
                '單位': '',
                '段落內容': '在此份ESG報告中未找到再生塑膠或永續材料相關的數值數據',
                '段落編號': '',
                '頁碼': '',
                '信心分數': 0.0,
                '上下文': '可能的原因：1) 該公司未涉及相關業務 2) 報告中未詳細披露相關數據 3) 關鍵字匹配範圍需要調整'
            }
            main_data.append(no_result_row)
        
        # 統計數據
        stats_data = []
        if len(extractions) > 0:
            for keyword, count in summary.keywords_found.items():
                keyword_extractions = [e for e in extractions if e.keyword == keyword]
                
                stats_data.append({
                    '關鍵字': keyword,
                    '提取數量': count,
                    '平均信心分數': round(np.mean([e.confidence for e in keyword_extractions]) if keyword_extractions else 0, 3),
                    '最高信心分數': round(max([e.confidence for e in keyword_extractions]) if keyword_extractions else 0, 3)
                })
        else:
            stats_data.append({
                '關鍵字': '搜尋摘要',
                '提取數量': 0,
                '平均信心分數': 0.0,
                '最高信心分數': 0.0
            })
        
        # 寫入Excel
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # 主要結果工作表
            sheet_name = '提取結果' if len(extractions) > 0 else '無提取結果'
            pd.DataFrame(main_data).to_excel(writer, sheet_name=sheet_name, index=False)
            
            # 統計工作表
            pd.DataFrame(stats_data).to_excel(writer, sheet_name='統計摘要', index=False)
            
            # 處理摘要（包含股票代號信息）
            summary_data = [{
                '股票代號': stock_code or 'N/A',
                '公司名稱': doc_info.company_name,
                '公司簡稱': short_company_name,
                '報告年度': doc_info.report_year,
                '總文檔數': summary.total_documents,
                '總提取結果': summary.total_extractions,
                '處理狀態': '成功提取' if len(extractions) > 0 else '無相關數據',
                '處理時間(秒)': round(summary.processing_time, 2),
                '處理日期': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                '提取器版本': 'ESG報告書提取器 v2.0 增強版',
                '向量資料庫': Path(doc_info.db_path).name
            }]
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='處理摘要', index=False)
        
        if len(extractions) > 0:
            print(f"✅ Excel檔案已保存，包含 {len(extractions)} 項提取結果")
        else:
            print(f"✅ Excel檔案已保存，標記為無提取結果")
        
        return output_path
    
    # 輔助方法（保持原有邏輯）
    def _split_paragraphs(self, text: str) -> List[str]:
        """段落分割"""
        paragraphs = []
        
        # 標準分割
        standard_paras = re.split(r'\n{2,}|\r{2,}', text)
        paragraphs.extend([p.strip() for p in standard_paras if len(p.strip()) >= 20])
        
        # 句號分割
        sentence_paras = re.split(r'。{2,}|\.{2,}', text)
        paragraphs.extend([p.strip() for p in sentence_paras if len(p.strip()) >= 40])
        
        # 保持原文
        if len(text.strip()) >= 60:
            paragraphs.append(text.strip())
        
        # 去重
        unique_paragraphs = []
        seen = set()
        for para in paragraphs:
            para_hash = hash(para[:100])
            if para_hash not in seen:
                seen.add(para_hash)
                unique_paragraphs.append(para)
        
        return unique_paragraphs
    
    def _extract_unit(self, value_str: str) -> str:
        """從數值字符串中提取單位"""
        units = re.findall(r'[a-zA-Z\u4e00-\u9fff]+', value_str)
        return units[-1] if units else ""
    
    def _get_context_window(self, full_text: str, target_paragraph: str, window_size: int = 200) -> str:
        """獲取段落的上下文窗口"""
        try:
            pos = full_text.find(target_paragraph)
            if pos == -1:
                return target_paragraph[:400]
            
            start = max(0, pos - window_size)
            end = min(len(full_text), pos + len(target_paragraph) + window_size)
            
            return full_text[start:end]
        except:
            return target_paragraph[:400]

# =============================================================================
# 為了向後兼容，保留原有類名的別名
# =============================================================================

ESGExtractor = EnhancedESGExtractor  # 別名，確保主程式可以正常使用

def main():
    """主函數 - 測試用"""
    print("📊 增強版ESG報告書提取器測試模式")
    
    extractor = EnhancedESGExtractor(enable_llm=False)
    print("✅ 增強版ESG報告書提取器初始化完成")

if __name__ == "__main__":
    main()